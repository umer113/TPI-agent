
from math import ceil
import os
import shutil
import subprocess
import asyncio
from datetime import datetime
from dotenv import load_dotenv
import pandas as pd
import streamlit as st
from openai import AsyncOpenAI
from groq import Groq
import json
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import csv
import io
import requests
import hashlib
import sys
from docx import Document  # New import for creating .docx files
from io import BytesIO  # For handling in-memory file

load_dotenv()

FIRST_PAGE_URL = {
    "DVA Minister":    "https://minister.dva.gov.au/minister-media-releases?page=1",
    "DVA Veteran Affairs": "https://www.dva.gov.au/about/news/vetaffairs",
    "DVA Repatriation Commission": "https://www.dva.gov.au/about/overview/repatriation-commission/gwen-cherne-veteran-family-advocate-commissioner/veteran-family-advocate-commissioner-gwen-cherne",
    "DVA Website About":"https://www.dva.gov.au/about/our-work-response-royal-commission-defence-and-veteran-suicid",
    "DVA Website Home":"https://clik.dva.gov.au/",
    "DVA Website Latest News":"https://www.dva.gov.au/about/news/latest-news"
}

API_KEY = os.getenv("OPEN_AI_API_KEY")
if not API_KEY:
    st.error("Please set the OPENAI_API_KEY in your .env file.")
    st.stop()

OPENAI_API_KEY = API_KEY
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

CHAT_DIR = "chat_history"
DATA_DIR = "data"
os.makedirs(CHAT_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)

# ——— Chat Persistence ———
def _clean_filename(s: str) -> str:
    return "".join(c if c.isalnum() or c in "-_ " else "_" for c in s).strip().replace(" ", "_")

def save_chat(history, chat_id=None, title=None):
    if title:
        new_id = _clean_filename(title)[:50] or datetime.now().strftime("%Y%m%d_%H%M%S")
    elif chat_id:
        new_id = chat_id
    else:
        new_id = datetime.now().strftime("%Y%m%d_%H%M%S")

    new_path = os.path.join(CHAT_DIR, f"{new_id}.json")
    if chat_id and chat_id != new_id:
        old_path = os.path.join(CHAT_DIR, f"{chat_id}.json")
        if os.path.exists(old_path):
            os.remove(old_path)

    with open(new_path, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)
    return new_id

def load_chats():
    chats = []
    for fname in sorted(os.listdir(CHAT_DIR), reverse=True):
        if not fname.endswith(".json"):
            continue
        path = os.path.join(CHAT_DIR, fname)
        with open(path, "r", encoding="utf-8") as f:
            msgs = json.load(f)
        chats.append({"id": fname[:-5], "messages": msgs})
    return chats

# ——— Scraper Runner ———
def list_scrapers(scraper_dir="scrapers"):
    return [f[:-3] for f in os.listdir(scraper_dir)
            if f.endswith(".py") and not f.startswith("__")]

common_headers = {
    "sec-ch-ua-platform": "Android",
    "User-Agent": (
        "Mozilla/5.0 (Linux; Android 6.0; Nexus 5 Build/MRA58N)"
        " AppleWebKit/537.36 (KHTML, like Gecko) Chrome/134.0.0.0"
        " Mobile Safari/537.36"
    ),
    "sec-ch-ua": '"Chromium";v="134", "Not:A-Brand";v="24", "Google Chrome";v="134"',
    "sec-ch-ua-mobile": "?1",
}

common_cookies = {
    "_gid": "GA1.3.1046864615.1743240670",
    "monsido": "4081743240673336",
    "_ga_XXXXXXXX": "GS1.1.1743243185.2.1.1743244433.0.0.0",
    "bm_mi": "3C2F88576EB2B1328DF4957982B85D2D~YAAQvvzUF1LjGqOVAQAAWGSd6Bth1EnnrqM/55ra+zZ0CT0o/"
             "5KLuglk/gQSB7kCoLjQwgCbOP906LWlWZpl4fyxcq+yuzGM8msirSFwu1nYdAotFYTHknHGqft33p+"
             "DMIqxmyzvdQzeuYdus7Xtt+oHgGiH8SCgPKX1NtMBWZW5lrG7FfXOfvaS8Odl3AA6lUi25CyUP+fK7"
             "uNQhboYal3H0DmCqbBPi5mqlDApqeGHtAMdQKrVixy2OwbwEhSMMuabDb2ibFZ+tu0ohB4YO1xQHwc"
             "FgoOG6YNswq0nSqtQBryENbhxkjofmazHpE8JywMoO2eWWQm3Txnd52nHkh6EaeI=~1",
    "_gat_gtag_UA_129191797_1": "1",
    "_gat_gtag_UA_67990327_1": "1",
    "_ga_MN793G4JHJ": "GS1.1.1743364571.3.0.1743364571.0.0.0",
    "_ga_FT6SLY9TWT": "GS1.1.1743364376.9.1.1743364574.0.0.0",
    "_ga_0XT7NFV9ZS": "GS1.1.1743364376.9.1.1743364574.0.0.0",
    "bm_sv": "AF7F1D971ACA5FD425CC7DC6D72B9CBC~YAAQvvzUF3PjGqOVAQAAJqGg6Buy7dRTKosyL4YNrqYTl"
             "oJ4Bouxg3EjnJ3fZ0HOiZaZW6nbfsodMC9h0XpffP79Cs0AxpmAR4zH0aL3GIeC4Rhi7ozMlQBhupO"
             "lz+hXJ55VeO7KgaJtW6ym4VjIN/7yh4uk68j3bp+0VK+4ZudN6dkpyRXhfBQXhrNWcT96qjllYRrY"
             "EZ6ZZbPI34HZcdPfFJ0xtuu1BJcV0TFWPeeBL7e3zGyCiwLzvkpECEXA~1",
    "_ga": "GA1.1.1075414505.1743240668",
}

import random
proxies_list = [
    'beqcfgqd:zx2ta8sl24bs@91.217.72.56:6785',
    'beqcfgqd:zx2ta8sl24bs@103.37.181.190:6846',
    'beqcfgqd:zx2ta8sl24bs@45.43.183.159:6471',
    'beqcfgqd:zx2ta8sl24bs@64.137.18.245:6439',
    'beqcfgqd:zx2ta8sl24bs@104.238.50.211:6757',
    'beqcfgqd:zx2ta8sl24bs@89.249.192.133:6532',
    'beqcfgqd:zx2ta8sl24bs@103.101.88.235:5959',
    'beqcfgqd:zx2ta8sl24bs@145.223.45.130:6983',
    'beqcfgqd:zx2ta8sl24bs@45.38.78.112:6049',
]

def fetch_page_with_proxy(
    url,
    proxies_list,
    headers=None,
    cookies=None,
    max_tries=5,
    timeout=10
):
    tried = set()
    attempts = min(max_tries, len(proxies_list))
    for _ in range(attempts):
        proxy = random.choice(proxies_list)
        if proxy in tried:
            continue
        tried.add(proxy)

        proxy_cfg = {
            "http":  f"http://{proxy}",
            "https": f"http://{proxy}"
        }

        try:
            resp = requests.get(
                url,
                headers=headers,
                cookies=cookies,
                proxies=proxy_cfg,
                timeout=timeout
            )
            if resp.status_code == 200:
                return resp
            else:
                continue
        except Exception:
            continue

    raise RuntimeError(f"All {len(tried)} proxies failed for {url}")

def get_top_n_listings(
    url,
    proxies,
    headers=None,
    cookies=None,
    n=5,
    max_tries=5,
    timeout=10
):
    resp = fetch_page_with_proxy(
        url,
        proxies_list=proxies,
        headers=headers,
        cookies=cookies,
        max_tries=max_tries,
        timeout=timeout
    )
    soup = BeautifulSoup(resp.text, "html.parser")
    anchors = soup.select("a.card")[:n]
    return [urljoin(url, a["href"]) for a in anchors if a.get("href")]

def run_scraper(module_name, scraper_dir="scrapers"):
    url = FIRST_PAGE_URL.get(module_name)
    meta_json = os.path.join(DATA_DIR, f"{module_name}_top5.json")
    top5 = None

    if url:
        try:
            top5 = get_top_n_listings(
                url,
                proxies=proxies_list,
                headers=common_headers,
                cookies=common_cookies,
                n=5
            )

            if os.path.exists(meta_json):
                with open(meta_json, "r") as f:
                    old_top5 = json.load(f)
                if old_top5 == top5:
                    st.sidebar.info("✅ No new updates detected. Your dataset is already current, so the scrape has been skipped.")
                    return

            with open(meta_json, "w") as f:
                json.dump(top5, f, indent=2)

        except Exception as e:
            st.sidebar.warning(f"Top-5 fetch failed: {e}\n→ running full scrape anyway.")

    before = set(os.listdir(DATA_DIR))
    script = os.path.join(scraper_dir, f"{module_name}.py")
    try:
        subprocess.run([sys.executable, script], check=True)
    except Exception as e:
        st.sidebar.error(f"Scraper error: {e}")
        return

    added = set(os.listdir(DATA_DIR)) - before
    if added:
        for fn in added:
            st.sidebar.success(f"New CSV: {fn}")
    else:
        st.sidebar.info("No CSV produced.")

import tiktoken

groq_model = "meta-llama/llama-4-scout-17b-16e-instruct"

async def ask_agent(csv_text: str, question: str, model: str, chat_history: list) -> str:
    use_groq = model.startswith("meta-llama/")
    try:
        encoding = tiktoken.encoding_for_model(model)
    except KeyError:
        encoding = tiktoken.get_encoding("p50k_base" if use_groq else "cl100k_base")

    def count_tokens(text: str) -> int:
        return len(encoding.encode(text))

    system_prompt = (
        "You are a data-savvy AI assistant. "
        "Using ONLY the provided CSV data and conversation history, respond clearly and accurately to the user’s question. "
        "- Reference relevant rows or columns explicitly. "
        "- Present findings in a structured format (e.g., bullet points, numbered lists, or short paragraphs). "
        "- Do not assume any data not in the CSV. "
        "- If information is missing, state it clearly."
    )

    history_context = "".join(
        f"{('User' if m['role']=='user' else 'Assistant')}: {m['content']}\n\n"
        for m in chat_history
    )

    MODEL_MAX = 16385
    HEADROOM = 512
    usable_tokens = MODEL_MAX - HEADROOM

    static_tokens = (
        count_tokens(system_prompt)
        + count_tokens("### Conversation History:\n")
        + count_tokens(history_context)
        + count_tokens("### User Question:\n")
        + count_tokens(question)
    )

    async def send_chat(prompt: str) -> str:
        client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

        resp = await client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ]
        )
        return resp.choices[0].message.content.strip()


    async def send_groq(prompt: str) -> str:
        def run_sync():
            client = Groq(api_key=os.getenv("GROQ_API_KEY"))
            resp = client.chat.completions.create(
                model=groq_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ]
            )
            return resp.choices[0].message.content.strip()
        return await asyncio.to_thread(run_sync)

    def make_prompt(csv_section: str) -> str:
        return (
            f"### Conversation History:\n{history_context}"
            f"### CSV Data:\n{csv_section}\n\n"
            f"### User Question:\n{question}"
        )

    full_prompt = make_prompt(csv_text)
    if static_tokens + count_tokens(full_prompt) <= usable_tokens:
        return await (send_groq(full_prompt) if use_groq else send_chat(full_prompt))

    lines = csv_text.split("\n")
    header, rows = lines[0], lines[1:]
    avg_tokens_per_row = max(1, count_tokens("\n".join(rows)) // len(rows))
    rows_per_chunk = max(1, (usable_tokens - static_tokens) // avg_tokens_per_row)

    while True:
        chunks = [rows[i:i+rows_per_chunk] for i in range(0, len(rows), rows_per_chunk)]
        if all(
            static_tokens + count_tokens(make_prompt(header + "\n" + "\n".join(chunk))) <= usable_tokens
            for chunk in chunks
        ):
            break
        rows_per_chunk = max(1, rows_per_chunk // 2)

    partials = []
    for chunk in chunks:
        prompt = make_prompt(header + "\n" + "\n".join(chunk))
        part = await (send_groq(prompt) if use_groq else send_chat(prompt))
        partials.append(part)

    synthesis = (
        "Please combine the following partial responses into a single, well-structured answer to the user's question:\n\n"
        + "\n---\n".join(partials)
    )
    return await (send_groq(synthesis) if use_groq else send_chat(synthesis))

# ——— Function to create .docx file ———
def create_docx(content: str) -> BytesIO:
    doc = Document()
    doc.add_heading("Assistant Response", level=1)
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.markdown("""
    <style>
      .user-message {
        border: 1px solid #6B7280;
        color: #F9FAFB;
        padding: 0.75rem 1rem;
        border-radius: 0.375rem;
        margin: 0.5rem 0;
        max-width: 80%;
      }
    </style>
    """, unsafe_allow_html=True)

    st.title("🕸️ TPI Overwatch AI")

    if "query" not in st.session_state:
        st.session_state["query"] = ""

    st.sidebar.image("logo.png", width=200)

    scrapers = list_scrapers()
    choice = st.sidebar.selectbox("Select Source", scrapers)

    raw_model = st.sidebar.selectbox(
        "Model",
        ["gpt-3.5-turbo-16k", "Groq"],
        key="model_select"
    )
    model = raw_model if raw_model != "Groq" else "meta-llama/llama-4-scout-17b-16e-instruct"


    if st.sidebar.button("Fetch Latest Content"):
        df = run_scraper(choice)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")

        if isinstance(df, pd.DataFrame):
            fn = f"{choice}_{ts}.csv"
            os.makedirs("data", exist_ok=True)
            df.to_csv(os.path.join("data", fn), index=False)
            st.sidebar.success(f"Saved {fn}")
        elif df is None:
            src = f"{choice}.csv"
            if os.path.exists(src):
                fn = f"{choice}_{ts}.csv"
                os.makedirs("data", exist_ok=True)
                shutil.move(src, os.path.join("data", fn))
                st.sidebar.success(f"Scraping done {src} → data/{fn}")
            else:
                st.sidebar.error(f"No DataFrame returned and `{src}` not found.")
        else:
            st.sidebar.error("Scraper did not return a DataFrame.")

    files = sorted(
        [f for f in os.listdir("data") if f.lower().endswith(".csv")],
        key=lambda f: os.path.getmtime(os.path.join("data", f)),
        reverse=True
    )
    if not files:
        st.sidebar.info("No data yet. Run a scraper.")
        return

    labels = [
        f"{f} — {datetime.fromtimestamp(os.path.getmtime(os.path.join('data', f))).strftime('%Y-%m-%d %H:%M:%S')}"
        for f in files
    ]
    sel = st.sidebar.selectbox("Focus", labels)
    sel_file = files[labels.index(sel)]
    df = pd.read_csv(os.path.join("data", sel_file))

    chats = load_chats()
    options = ["🆕 New Article Thread"] + [
        c.get("title", c['id']) for c in chats
    ]
    default_idx = 0
    if st.session_state.get("chat_id"):
        default_idx = next(
            (i+1 for i, c in enumerate(chats) if c["id"] == st.session_state.chat_id),
            0
        )
    sel = st.sidebar.selectbox(
        "Article Threads",
        options,
        index=default_idx,
        key="chat_select"
    )
    if sel == "🆕 New Article Thread":
        st.session_state.chat_id = None
        st.session_state.chat_history = []
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        existing_titles = {c.get("title", "") for c in chats}
        unique_ts = ts
        counter = 1
        while unique_ts in existing_titles:
            unique_ts = f"{ts}_{counter}"
            counter += 1
        st.session_state["new_chat_title"] = unique_ts
    else:
        idx = options.index(sel) - 1
        st.session_state.chat_id = chats[idx]["id"]
        st.session_state.chat_history = chats[idx]["messages"].copy()

    # raw_model = st.sidebar.selectbox(
    #     "Model",
    #     ["gpt-3.5-turbo-16k", "Groq"],
    #     key="model_select"
    # )
    # model = raw_model if raw_model != "Groq" else "meta-llama/llama-4-scout-17b-16e-instruct"

    st.subheader(f"Dataset: {sel_file}")
    st.dataframe(df)
    st.markdown("---")

    st.sidebar.header("Predefined Prompts")
    predef_prompts = [
        "Summarize the key insights from this dataset.",
        "Write a comprehensive article based on this dataset, weaving in key insights, context, and potential implications.",
        "Give me a narrative overview of what this data represents."
    ]
    def _set_query(p):
        st.session_state["query"] = p

    for i, p in enumerate(predef_prompts):
        st.sidebar.button(
            p,
            key=f"predef_{i}",
            on_click=_set_query,
            args=(p,)
        )

    with st.form("chat_form", clear_on_submit=False):
        query = st.text_input("Ask anything—article, summary, insight…", key="query")
        submitted = st.form_submit_button("Ask Agent")
        if submitted and query:
            st.session_state.chat_history.append({"role": "user", "content": query})
            csv_text = df.to_csv(index=False)
            with st.spinner("🤖 Agent is thinking..."):
                answer = asyncio.run(ask_agent(csv_text, query, model, st.session_state.chat_history))
            st.session_state.chat_history.append({"role": "assistant", "content": answer})
            csv_basename = os.path.splitext(sel_file)[0]
            chat_title = st.session_state.get("new_chat_title", "untitled")

            new_id = save_chat(
                st.session_state.chat_history,
                chat_id=st.session_state.get("chat_id"),
                title=chat_title
            )
            st.session_state.chat_id = new_id
            try:
                st.experimental_rerun()
            except AttributeError:
                st.rerun()

    # ——— Display Chat with Download Buttons ———
    for i, msg in enumerate(st.session_state.chat_history):
        if msg["role"] == "user":
            st.markdown(f'<div class="user-message">👤 {msg["content"]}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f"**🤖** {msg['content']}")
            # Create .docx file for the assistant's response
            docx_buffer = create_docx(msg["content"])
            st.download_button(
                label="Save Article Draft",
                data=docx_buffer,
                file_name=f"Article_{i}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{i}"
            )

    st.markdown("---")

if __name__ == "__main__":
    main()
